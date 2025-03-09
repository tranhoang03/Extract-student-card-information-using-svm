
import streamlit as st
import cv2
import os
import io
from image_processor import ImageProcessor, CoordinateLoader
from model_predictor import ModelPredictor
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
from PIL import Image
# Định nghĩa đường dẫn mô hình
MODEL_PATHS = {
    'hoten': 'https://drive.google.com/uc?export=download&id=1deVSjp9IVEV2tQd4HU2Sg8M5HYLRVnnR',
    'ngaysinh': 'https://drive.google.com/uc?export=download&id=1e3wk3RmALDowN8-qlril8ISheCKiLQPo',
    'lop': 'https://drive.google.com/uc?export=download&id=1XUB6HF55ANed461uYPyU4jwLwZOnOilw',
    'msv': 'https://drive.google.com/uc?export=download&id=1oahCLZWJyeN5fDSxzSs4KQiEXR39SFKr',  
    'nienkhoa': 'https://drive.google.com/uc?export=download&id=19YrPx9RtUo3nLOjLzB-kMNbzwCao6Pag',
    'anhthe': 'https://drive.google.com/uc?export=download&id=1JkLcsoxMfU-2zF8wNr61L_Hq3fVQu5uo'
}

# Tải các thông tin tọa độ từ XML
coordinate_loader = CoordinateLoader()
average_coordinates, max_hoten_box = coordinate_loader.load_coordinates_from_xml(r'training_data_segmentation/annotations.xml')
all_coordinates = coordinate_loader.get_all_coordinates(average_coordinates, max_hoten_box)

# Khởi tạo predictor
predictor = ModelPredictor(MODEL_PATHS)

# Hàm lưu thông tin sinh viên vào file Word
def save_student_info_to_word(all_predictions, all_extracted_info):
    doc = Document()
    doc.add_heading("Thông tin Sinh viên", level=1)

    for idx, (predictions, extracted_info) in enumerate(zip(all_predictions, all_extracted_info), start=1):
        doc.add_heading(f"Thông tin Thẻ {idx}", level=2)
        doc.add_heading("Thông tin chi tiết:", level=3)
        for label, predicted_class in predictions.items():
            if label != "anhthe":
                doc.add_paragraph(f"{label.capitalize()}: {predicted_class}")

        if "anhthe" in extracted_info:
            doc.add_heading("Ảnh thẻ:", level=3)
            temp_image_path = "temp_anhthe.jpg"
            cv2.imwrite(temp_image_path, extracted_info["anhthe"])
            doc.add_picture(temp_image_path, width=Inches(2))
            os.remove(temp_image_path)

        doc.add_page_break()

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Giao diện Streamlit
st.set_page_config(page_title="Trích xuất Thông tin Thẻ Sinh viên", layout="wide")
st.markdown(
    """
    <style>
    .main {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
    }
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        font-size: 16px;
        border-radius: 8px;
        height: 50px;
    }
    .stMarkdown h1 {
        color: #30336b;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("🎓 Trích xuất thông tin thẻ sinh viên")
st.markdown("**Ứng dụng giúp trích xuất thông tin từ ảnh thẻ sinh viên và lưu lại vào file Word.**")
st.markdown("---")

uploaded_files = st.file_uploader("📂 Tải lên ảnh thẻ sinh viên", type=["jpg", "png"], accept_multiple_files=True)

if st.button("🖼️ Sử dụng ảnh mẫu"):
    default_image_url = "https://raw.githubusercontent.com/tranhoang03/Extract-student-card-information-using-svm/master/z5424934443026_83aff27331bd9d2087ed8bbf7b11120c.jpg"
    response = requests.get(default_image_url)
    response.raise_for_status()
    uploaded_files = [Image.open(BytesIO(response.content))]
    uploaded_files[0].name = "default_image.jpg"
if uploaded_files:
    all_predictions = []
    all_extracted_info = []

    with st.spinner("⏳ Đang xử lý hình ảnh..."):
        for uploaded_file in uploaded_files:
            image_path = uploaded_file.name
            with open(image_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.image(image_path, caption=f"Ảnh gốc - {uploaded_file.name}", use_column_width=True)

            processor = ImageProcessor()
            cropped_card = processor.crop_card(image_path)

            if cropped_card is not None:
                extracted_info = processor.crop_info_from_coordinates(cropped_card, all_coordinates)
                predictions = predictor.predict_info(extracted_info)

                st.markdown("### 🔍 Kết quả dự đoán:")
                for label, pred in predictions.items():
                    st.write(f"**{label.capitalize()}**: {pred}")

                st.markdown("### ✂️ Các vùng đã tách:")
                cols = st.columns(len(extracted_info))
                for idx, (label, img) in enumerate(extracted_info.items()):
                    with cols[idx]:
                        st.image(cv2.cvtColor(img, cv2.COLOR_BGR2RGB), caption=label)

                all_predictions.append(predictions)
                all_extracted_info.append(extracted_info)

            os.remove(image_path)

    if all_predictions:
        byte_io = save_student_info_to_word(all_predictions, all_extracted_info)
        st.download_button("💾 Tải về file word", byte_io, file_name="thong_tin_sinh_vien.docx")
